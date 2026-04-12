""""
OUTPUT:
    DTR_<Name>_<Month>_<Year>.docx  — printable Word document
    DTR_<Name>_<Month>_<Year>.pdf   — if LibreOffice is installed

DAY TYPES:
    Regular          — manual AM/PM time in and out
    Weekend          — auto-skipped (Saturday & Sunday)
    Holiday          — add to HOLIDAYS dict below
    Travel Order (TO)— auto 8hrs, enter reference number
    Travel Ticket (TT)— auto 8hrs, enter reference number
    Absent           — marked, no hours counted
    Missing Log (ML) — dashes shown, flagged for follow-up
"""

import os
import calendar
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================
#  CONFIGURATION — edit this section before running
# =============================================================

OFFICE_LINE1 = "Republic of the Philippines"
OFFICE_LINE2 = "Bureau of Fisheries and Aquatic Resources"
OFFICE_LINE3 = "Regional Office No. III"
DOC_TITLE    = "DAILY TIME RECORD"

# ── Output folder — change this to wherever you want files saved
OUTPUT_DIR = r"C:\Users\daves\Downloads\MY_PROJECTS\DTR_GENERATOR\DTR_TEMPLATES"

EXPECTED_AM_IN  = "08:00"
EXPECTED_AM_OUT = "12:00"
EXPECTED_PM_IN  = "13:00"
EXPECTED_PM_OUT = "17:00"
AUTO_HOURS      = 8.0   # hours credited for TO / TT days

# ── Add holidays here: "YYYY-MM-DD": "Holiday Label"
# ── The script will automatically mark these days as holidays
HOLIDAYS = {
    "2026-02-04": "Special Non-Working Holiday",
    "2026-02-17": "Special Non-Working Holiday",
    "2026-02-25": "People Power Revolution Anniversary",
    # add more as needed:
    # "2026-04-09": "Araw ng Kagitingan",
    # "2026-06-12": "Independence Day",
}


# =============================================================
#  TIME HELPERS
# =============================================================

def to_min(t: str) -> int:
    if not t or not t.strip():
        return 0
    h, m = map(int, t.strip().split(":"))
    return h * 60 + m

# ── Individual deduction computers ──────────────────────────

def compute_late(am_in: str) -> int:
    """Minutes late — AM In after 08:00."""
    if not am_in:
        return 0
    return max(0, to_min(am_in) - to_min(EXPECTED_AM_IN))

def compute_early_lunch(am_out: str) -> int:
    """Minutes early lunch — AM Out before 12:00."""
    if not am_out:
        return 0
    return max(0, to_min(EXPECTED_AM_OUT) - to_min(am_out))

def compute_late_lunch(pm_in: str) -> int:
    """Minutes late return from lunch — PM In after 13:00."""
    if not pm_in:
        return 0
    return max(0, to_min(pm_in) - to_min(EXPECTED_PM_IN))

def compute_undertime(pm_out: str) -> int:
    """Minutes undertime — PM Out before 17:00."""
    if not pm_out:
        return 0
    return max(0, to_min(EXPECTED_PM_OUT) - to_min(pm_out))

# ── Missing entry deduction rules ───────────────────────────
# Rule: 1 missing entry = 4 hrs deduction
#       AM Out + PM In both missing (only 2 logs) = 4 hrs (not 8)
#       3 or more missing = full day

def compute_missing_deduction(am_in, am_out, pm_in, pm_out) -> int:
    """
    Returns deduction in minutes based on missing entries.
    Missing AM Out only       → 240 min (4 hrs)
    Missing PM In only        → 240 min (4 hrs)
    Missing AM Out + PM In    → 240 min (4 hrs) — treated as one lunch block
    Missing 3 entries         → 480 min (full day = 8 hrs)
    Missing all 4             → 480 min (full day)
    """
    missing = [
        not am_in,
        not am_out,
        not pm_in,
        not pm_out,
    ]
    count = sum(missing)
    if count == 0:
        return 0
    if count >= 3:
        return 480  # full day
    # AM Out and PM In both missing (lunch block) = 4 hrs, not 8
    if not am_out and not pm_in:
        return 240
    # Single missing entry = 4 hrs
    if count == 1:
        return 240
    # 2 missing but not the lunch pair = 8 hrs (full day)
    return 480

def compute_hours(am_in, am_out, pm_in, pm_out) -> float:
    """Hours worked from available entries."""
    total = 0
    try:
        if am_in and am_out:
            total += max(0, to_min(am_out) - to_min(am_in))
        if pm_in and pm_out:
            total += max(0, to_min(pm_out) - to_min(pm_in))
        # Only 2 logs: AM In + PM Out — compute full span minus 1hr lunch
        if am_in and pm_out and not am_out and not pm_in:
            span = to_min(pm_out) - to_min(am_in)
            total = max(0, span - 60)  # subtract 1 hr assumed lunch
    except Exception:
        return 0.0
    return round(total / 60, 2)

def compute_all_deductions(am_in, am_out, pm_in, pm_out):
    """
    Returns dict of all deductions in minutes.
    Missing entries take priority over time-based deductions.
    """
    miss_ded = compute_missing_deduction(am_in, am_out, pm_in, pm_out)

    # Only compute time deductions if the entry exists
    late        = compute_late(am_in)         if am_in  else 0
    early_lunch = compute_early_lunch(am_out) if am_out else 0
    late_lunch  = compute_late_lunch(pm_in)   if pm_in  else 0
    undertime   = compute_undertime(pm_out)   if pm_out else 0

    return {
        "late":        late,
        "early_lunch": early_lunch,
        "late_lunch":  late_lunch,
        "undertime":   undertime,
        "missing_ded": miss_ded,
        "total":       late + early_lunch + late_lunch + undertime + miss_ded,
    }

def mins_to_readable(mins: int) -> str:
    if mins <= 0:
        return "none"
    h = mins // 60
    m = mins % 60
    if h and m:
        return f"{h} hr {m} min"
    return f"{h} hr" if h else f"{m} min"

def mins_to_days_hrs_mins(mins: int) -> str:
    """Convert minutes to working days, hours, minutes.
    1 working day = 480 minutes (8 hours).
    e.g. 505 min → 1 day  0 hrs  25 min
    """
    if mins <= 0:
        return "0 days  0 hrs  0 min"
    MINS_PER_DAY = 480
    days = mins // MINS_PER_DAY
    rem  = mins % MINS_PER_DAY
    hrs  = rem // 60
    mn   = rem % 60
    return f"{days} day{'s' if days != 1 else ''}  {hrs} hr{'s' if hrs != 1 else ''}  {mn} min"


# =============================================================
#  INPUT HELPERS
# =============================================================

def ask(label: str, default: str = "") -> str:
    raw = input(f"  {label}{f' [{default}]' if default else ''}: ").strip()
    return raw if raw else default

def ask_int(label: str, default: int) -> int:
    while True:
        raw = input(f"  {label} [{default}]: ").strip()
        if not raw:
            return default
        try:
            return int(raw)
        except ValueError:
            print("    Please enter a number.")

def ask_time(label: str, default: str = "") -> str:
    while True:
        hint = default if default else "leave blank to skip"
        raw = input(f"    {label} [{hint}]: ").strip()
        if not raw:
            return default if default else ""
        try:
            datetime.strptime(raw, "%H:%M")
            return raw
        except ValueError:
            print("      Format must be HH:MM — e.g. 08:05 or 17:00")


# =============================================================
#  COLLECT EMPLOYEE INFO
# =============================================================

def collect_info() -> dict:
    print()
    print("=" * 60)
    print("  BFAR-III Daily Time Record Generator")
    print("=" * 60)
    print("\n  [EMPLOYEE INFORMATION]")
    info = {
        "id_no":    ask("ID No.").upper(),
        "name":     ask("Full Name  (e.g. DELA CRUZ, JUAN A.)").upper(),
        "position": ask("Position / Designation").upper(),
        "dept":     ask("Department / Section").upper(),
    }
    print("\n  [PERIOD COVERED]")
    info["year"]          = ask_int("Year", datetime.now().year)
    info["month"]         = ask_int("Month (1-12)", datetime.now().month)
    info["month_name"]    = calendar.month_name[info["month"]].upper()
    info["days_in_month"] = calendar.monthrange(info["year"], info["month"])[1]
    return info


# =============================================================
#  COLLECT DAILY ATTENDANCE
# =============================================================

def collect_attendance(info: dict) -> list:
    year  = info["year"]
    month = info["month"]
    days  = info["days_in_month"]

    print(f"\n  [DAILY ATTENDANCE — {info['month_name']} {year}]")
    print("  Press ENTER to accept default time.")
    print("  Leave blank to skip that entry.\n")

    records = []

    for day_num in range(1, days + 1):
        d        = date(year, month, day_num)
        day_name = d.strftime("%A").upper()
        date_str = d.strftime("%Y-%m-%d")
        is_we    = d.weekday() >= 5

        # ── Weekend — skip automatically
        if is_we:
            print(f"  [{date_str}] {day_name} — weekend, skipped")
            records.append(_make_rec(date_str, day_name,
                                     is_weekend=True))
            continue

        # ── Holiday — check HOLIDAYS dict
        if date_str in HOLIDAYS:
            label = HOLIDAYS[date_str]
            print(f"  [{date_str}] HOLIDAY — {label}")
            records.append(_make_rec(date_str, day_name,
                                     is_holiday=True,
                                     holiday_label=label))
            continue

        # ── Working day — ask what type
        print(f"\n  [{date_str}  {day_name}]")
        print("    Type:  [ENTER] Regular   [T] Travel Order/Ticket   [A] Absent   [M] Missing Log")
        day_type = input("    > ").strip().upper()

        # ── Travel Order or Travel Ticket
        if day_type == "T":
            kind = input("    TO or TT? [TO/TT]: ").strip().upper()
            kind = kind if kind in ("TO", "TT") else "TO"
            ref  = input(f"    {kind} Reference No.: ").strip()
            label = f"{kind}# {ref}" if ref else kind
            print(f"    → {label}  —  credited 8 hours, auto present")
            records.append(_make_rec(date_str, day_name,
                                     is_travel=True,
                                     travel_label=label,
                                     hours=AUTO_HOURS))
            continue

        # ── Absent
        if day_type == "A":
            print("    → Marked absent")
            records.append(_make_rec(date_str, day_name, is_absent=True))
            continue

        # ── Missing Log
        if day_type == "M":
            print("    → Missing log — will show dashes, flagged for follow-up")
            records.append(_make_rec(date_str, day_name, is_missing=True))
            continue

        # ── Regular day — enter times
        am_in  = ask_time("AM Time In  ", EXPECTED_AM_IN)
        am_out = ask_time("AM Time Out ", EXPECTED_AM_OUT)
        pm_in  = ask_time("PM Time In  ", EXPECTED_PM_IN)
        pm_out = ask_time("PM Time Out ", EXPECTED_PM_OUT)

        ded   = compute_all_deductions(am_in, am_out, pm_in, pm_out)
        hours = compute_hours(am_in, am_out, pm_in, pm_out)

        # Show deduction breakdown in terminal
        if ded["late"]:
            print(f"    → Late AM In:        {ded['late']} min")
        if ded["early_lunch"]:
            print(f"    → Early Lunch Out:   {ded['early_lunch']} min")
        if ded["late_lunch"]:
            print(f"    → Late Lunch In:     {ded['late_lunch']} min")
        if ded["undertime"]:
            print(f"    → Undertime PM Out:  {ded['undertime']} min")
        if ded["missing_ded"]:
            print(f"    → Missing entry ded: {ded['missing_ded']} min")
        if ded["total"]:
            print(f"    → Total deduction:   {ded['total']} min  ({mins_to_readable(ded['total'])})")
        print(f"    → Hours worked:      {hours}")

        records.append(_make_rec(
            date_str, day_name,
            am_in=am_in, am_out=am_out,
            pm_in=pm_in, pm_out=pm_out,
            hours=hours,
            late=ded["late"],
            early_lunch=ded["early_lunch"],
            late_lunch=ded["late_lunch"],
            undertime=ded["undertime"],
            missing_ded=ded["missing_ded"],
            total_ded=ded["total"],
        ))

    return records


def _make_rec(date_str, day_name,
              am_in="", am_out="", pm_in="", pm_out="",
              hours=0.0,
              late=0, early_lunch=0, late_lunch=0,
              undertime=0, missing_ded=0, total_ded=0,
              is_weekend=False, is_holiday=False, holiday_label="",
              is_travel=False, travel_label="",
              is_absent=False, is_missing=False) -> dict:
    return {
        "date": date_str, "day_name": day_name,
        "am_in": am_in, "am_out": am_out,
        "pm_in": pm_in, "pm_out": pm_out,
        "hours":       hours,
        "late":        late,
        "early_lunch": early_lunch,
        "late_lunch":  late_lunch,
        "undertime":   undertime,
        "missing_ded": missing_ded,
        "total_ded":   total_ded,
        "is_weekend":  is_weekend,
        "is_holiday":  is_holiday, "holiday_label": holiday_label,
        "is_travel":   is_travel,  "travel_label":  travel_label,
        "is_absent":   is_absent,
        "is_missing":  is_missing,
    }


# =============================================================
#  DOCX HELPERS
# =============================================================

def _shade(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _margins(cell, top=50, bot=50, left=80, right=80):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for s, v in [('top', top), ('bottom', bot),
                 ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:w'), str(v))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def _write(cell, text, bold=False, size=8.5, italic=False,
           align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if str(text) == "":
        return
    r = p.add_run(str(text))
    r.font.name   = "Arial"
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)

def _para(doc, text="", bold=False, size=10, italic=False,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          sa=2, sb=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    if text:
        r = p.add_run(text)
        r.font.name   = "Arial"
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


# =============================================================
#  BUILD DOCX
# =============================================================

def build_docx(info: dict, records: list, out_path: str,
               paper: str = "legal"):
    """
    paper = "legal"  → 8.5 x 13 inches  (default, recommended)
    paper = "short"  → 8.5 x 11 inches  (short bond / letter)
    Both fit on a single page.
    """
    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(13) if paper == "legal" else Inches(11)

    # Margins — slightly tighter on short bond to fit everything
    if paper == "legal":
        sec.top_margin = sec.bottom_margin = Inches(0.4)
        sec.left_margin = sec.right_margin = Inches(0.6)
    else:
        sec.top_margin = sec.bottom_margin = Inches(0.3)
        sec.left_margin = sec.right_margin = Inches(0.55)

    # ── Color palette
    NAVY    = "1A3A5C"   # header bar, column headers
    SILVER  = "D6DCE4"   # label cells in info block
    ALT     = "EBF3FB"   # alternating row tint
    WE_BG   = "F4F4F4"   # weekend rows
    HOL_BG  = "FFF8E7"   # holiday rows
    TRV_BG  = "EAF4EA"   # travel order/ticket rows
    ABS_BG  = "FDE9E9"   # absent rows
    MIS_BG  = "F9F9F9"   # missing log rows (light — B&W safe)
    TOT_BG  = "1A3A5C"   # totals header bar
    TOT_ROW = "EBF3FB"   # totals data row

    # ─────────────────────────────────────────
    #  FORMAL HEADER — compact spacing
    # ─────────────────────────────────────────
    _para(doc, OFFICE_LINE1, size=8,  sa=0)
    _para(doc, OFFICE_LINE2, size=9,  bold=True, sa=0)
    _para(doc, OFFICE_LINE3, size=8,  sa=2)

    # Thin divider line
    div_p = _para(doc, "", sa=2)
    pPr   = div_p._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:color'), NAVY)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Document title bar — reduced padding
    tt = doc.add_table(rows=1, cols=1)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc = tt.rows[0].cells[0]
    tc.width = Inches(7.3)
    _shade(tc, NAVY)
    _margins(tc, 70, 70, 140, 140)
    _write(tc, DOC_TITLE, bold=True, size=13, color=(255, 255, 255))
    _para(doc, "", sa=3)

    # ─────────────────────────────────────────
    #  EMPLOYEE INFO BLOCK — 2 rows
    #  Row 1: ID No. | value  | Month/Year | value
    #  Row 2: Name   | value  | Position/Dept | value
    # ─────────────────────────────────────────
    it = doc.add_table(rows=2, cols=4)
    it.style     = 'Table Grid'
    it.alignment = WD_TABLE_ALIGNMENT.CENTER

    IW = [Inches(0.9), Inches(1.8), Inches(1.1), Inches(3.5)]

    rows_data = [
        ("ID No.:",        info["id_no"],
         "Month / Year:",  f"{info['month_name']} {info['year']}"),
        ("Name:",          info["name"],
         "Position / Dept:", f"{info['position']}  |  {info['dept']}"),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(rows_data):
        for ci, (txt, is_lbl, w) in enumerate([
            (l1, True,  IW[0]), (v1, False, IW[1]),
            (l2, True,  IW[2]), (v2, False, IW[3]),
        ]):
            c = it.rows[ri].cells[ci]
            c.width = w
            if is_lbl:
                _shade(c, SILVER)
            _margins(c, 30, 30, 80, 50)
            _write(c, txt, bold=is_lbl, size=8.5,
                   align=WD_ALIGN_PARAGRAPH.LEFT)

    _para(doc, "", sa=3)

    # ─────────────────────────────────────────
    #  ATTENDANCE TABLE
    #  Cols: DATE | DAY | AM IN | AM OUT | PM IN | PM OUT | HRS | DEDUCTION
    # ─────────────────────────────────────────
    COL_W = [
        Inches(1.00),  # 0  DATE
        Inches(0.85),  # 1  DAY
        Inches(0.78),  # 2  AM IN
        Inches(0.78),  # 3  AM OUT
        Inches(0.78),  # 4  PM IN
        Inches(0.78),  # 5  PM OUT
        Inches(0.68),  # 6  HRS
        Inches(1.78),  # 7  DEDUCTION (combined, wider)
    ]
    HDRS = ["DATE", "DAY", "AM IN", "AM OUT",
            "PM IN", "PM OUT", "HRS", "DEDUCTION (min)"]
    NC   = len(COL_W)
    NR   = 1 + len(records) + 2

    # Row height — fixed for ALL rows (regular, absent, weekend, holiday, etc.)
    ROW_H = 280  # twips ≈ 0.19 inch — slightly taller to fill page nicely

    def _set_row_height(row):
        tr   = row._tr
        trPr = tr.get_or_add_trPr()
        trH  = OxmlElement('w:trHeight')
        trH.set(qn('w:val'),   str(ROW_H))
        trH.set(qn('w:hRule'), 'exact')   # exact = never grows, never shrinks
        trPr.append(trH)

    tbl = doc.add_table(rows=NR, cols=NC)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column header row — slightly taller than data rows
    hdr_row = tbl.rows[0]
    tr = hdr_row._tr; trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '300'); trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)
    for ci, (h, w) in enumerate(zip(HDRS, COL_W)):
        c = hdr_row.cells[ci]; c.width = w
        _shade(c, NAVY); _margins(c, 50, 50, 50, 50)
        _write(c, h, bold=True, size=8, color=(255, 255, 255))

    # ── Data rows
    total_ded_all = 0
    total_hrs     = 0.0
    days_present  = 0
    days_absent   = 0
    days_holiday  = 0

    for ri, rec in enumerate(records):
        row    = tbl.rows[1 + ri]
        _set_row_height(row)          # ← enforce uniform height every row
        is_we  = rec["is_weekend"]
        is_hol = rec["is_holiday"]
        is_trv = rec["is_travel"]
        is_abs = rec["is_absent"]
        is_mis = rec.get("is_missing", False)
        is_alt = (ri % 2 == 1) and not any([is_we, is_hol, is_trv, is_abs, is_mis])

        if is_we:    bg = WE_BG
        elif is_hol: bg = HOL_BG
        elif is_trv: bg = TRV_BG
        elif is_abs: bg = ABS_BG
        elif is_mis: bg = MIS_BG
        elif is_alt: bg = ALT
        else:        bg = "FFFFFF"

        # DATE cell
        dc = row.cells[0]; dc.width = COL_W[0]
        _shade(dc, bg); _margins(dc, 30, 30, 80, 40)
        _write(dc, rec["date"], size=8,
               color=(150, 150, 150) if is_we else None,
               align=WD_ALIGN_PARAGRAPH.LEFT)

        # DAY cell
        dyc = row.cells[1]; dyc.width = COL_W[1]
        _shade(dyc, bg); _margins(dyc, 20, 20, 60, 40)
        day_color = ((150, 150, 150) if is_we  else
                     (160, 110,   0) if is_hol else
                     ( 30, 120,  60) if is_trv else
                     (180,  60,  60) if is_abs else
                     (100, 100, 100) if is_mis else None)
        _write(dyc, rec["day_name"][:3], bold=True, size=7.5, color=day_color)

        # ── Holiday
        if is_hol:
            hol_words = rec['holiday_label'].upper().split()
            per_cell  = [""] * (NC - 2)
            for i, word in enumerate(hol_words):
                per_cell[i % len(per_cell)] += (" " if per_cell[i % len(per_cell)] else "") + word
            for ci in range(2, NC):
                c = row.cells[ci]; c.width = COL_W[ci]
                _shade(c, HOL_BG); _margins(c, 20, 20, 20, 20)
                _write(c, per_cell[ci - 2], bold=True, size=7,
                       italic=True, color=(160, 110, 0))
            days_holiday += 1
            continue

        # ── Weekend
        if is_we:
            for ci in range(2, NC):
                c = row.cells[ci]; c.width = COL_W[ci]
                _shade(c, WE_BG); _margins(c, 20, 20, 20, 20)
                _write(c, "", size=7.5)
            continue

        # ── Travel Order / Ticket
        if is_trv:
            label_parts = rec['travel_label'].split()
            per_cell    = [""] * (NC - 2)
            for i, word in enumerate(label_parts):
                per_cell[i % len(per_cell)] += (" " if per_cell[i % len(per_cell)] else "") + word
            for ci in range(2, NC):
                c = row.cells[ci]; c.width = COL_W[ci]
                _shade(c, TRV_BG); _margins(c, 20, 20, 20, 20)
                _write(c, per_cell[ci - 2], bold=True, size=7,
                       color=(30, 120, 60))
            total_hrs    += AUTO_HOURS
            days_present += 1
            continue

        # ── Absent
        if is_abs:
            absent_fill = ["", "", "", "ABSENT", "", ""]
            for ci in range(2, NC):
                c = row.cells[ci]; c.width = COL_W[ci]
                _shade(c, ABS_BG); _margins(c, 20, 20, 20, 20)
                _write(c, absent_fill[ci - 2], bold=True, size=7.5,
                       color=(180, 60, 60))
            days_absent += 1
            continue

        # ── Missing Log — no entries at all
        if is_mis:
            for ci in range(2, NC):
                c = row.cells[ci]; c.width = COL_W[ci]
                _shade(c, MIS_BG); _margins(c, 20, 20, 20, 20)
                _write(c, "———", size=7, color=(160, 160, 160))
            days_absent += 1
            continue

        # ── Regular day — fill time cells
        # Bold the time if it caused a deduction
        time_vals   = [rec["am_in"], rec["am_out"], rec["pm_in"], rec["pm_out"]]
        is_offender = [
            rec.get("late", 0)        > 0,  # AM IN  late
            rec.get("early_lunch", 0) > 0,  # AM OUT early
            rec.get("late_lunch", 0)  > 0,  # PM IN  late return
            rec.get("undertime", 0)   > 0,  # PM OUT early leave
        ]
        for ci, (v, w, offend) in enumerate(zip(time_vals, COL_W[2:6], is_offender)):
            c = row.cells[ci + 2]; c.width = w
            _shade(c, bg); _margins(c, 20, 20, 40, 40)
            is_blank = (v == "")
            _write(c, v if not is_blank else "?",
                   size=7.5,
                   bold=offend,
                   color=(180, 60, 60) if offend else ((200, 100, 0) if is_blank else None))

        # Hours worked
        hc = row.cells[6]; hc.width = COL_W[6]
        _shade(hc, bg); _margins(hc, 20, 20, 40, 40)
        _write(hc, str(rec["hours"]) if rec["hours"] > 0 else "", size=7.5)

        # Deduction cell — plain readable, no codes
        dc2 = row.cells[7]; dc2.width = COL_W[7]
        _shade(dc2, bg); _margins(dc2, 20, 20, 40, 40)
        ded = rec.get("total_ded", 0)
        if ded > 0:
            miss = rec.get("missing_ded", 0)
            time_ded = ded - miss
            # If 3+ entries missing → full day deduction
            if miss >= 480:
                ded_str = "1 day"
            elif ded >= 480:
                # full day from combined deductions
                d = ded // 480
                leftover = ded % 480
                h = leftover // 60
                m = leftover % 60
                if d and h and m:
                    ded_str = f"{d}d {h}h {m}m"
                elif d and h:
                    ded_str = f"{d}d {h}h"
                elif d and m:
                    ded_str = f"{d}d {m}m"
                else:
                    ded_str = f"{d}d"
            else:
                # Just hours and minutes
                h = ded // 60
                m = ded % 60
                if h and m:
                    ded_str = f"{h}h {m}m"
                elif h:
                    ded_str = f"{h}h"
                else:
                    ded_str = f"{m}m"
            _write(dc2, ded_str, bold=True, size=7.5, color=(180, 60, 60))
        else:
            _write(dc2, "", size=7.5)

        # Accumulate totals
        if rec["am_in"] or rec.get("total_ded", 0) > 0:
            total_ded_all += rec.get("total_ded", 0)
            total_hrs     += rec["hours"]
            days_present  += 1
        else:
            days_absent += 1

    # ─────────────────────────────────────────
    #  TOTALS — header bar + values row
    # ─────────────────────────────────────────
    base = 1 + len(records)

    lbar = tbl.rows[base]
    mc   = lbar.cells[0]
    for ci in range(1, NC):
        mc = mc.merge(lbar.cells[ci])
    _shade(mc, TOT_BG)
    _margins(mc, 50, 50, 100, 60)
    _write(mc, "MONTHLY TOTALS", bold=True, size=9,
           color=(255, 255, 255))

    # Deduction shown as: X days  X hrs  X min
    deduct_str = (
        f"Total Deduction: {mins_to_days_hrs_mins(total_ded_all)}"
    )
    vrow   = tbl.rows[base + 1]
    groups = [(0, 1), (2, 4), (5, 7)]
    labels = [
        f"Present: {days_present}   Absent: {days_absent}   Holiday: {days_holiday}",
        f"Total Hours Worked: {round(total_hrs, 2)} hrs",
        deduct_str,
    ]
    colors = [
        (20, 100, 60),
        (30, 80, 140),
        (180, 60, 60) if total_ded_all > 0 else (80, 80, 80),
    ]
    prev_end = -1
    for gi, (start, end) in enumerate(groups):
        if start <= prev_end:
            continue
        mc2 = vrow.cells[start]
        for ci in range(start + 1, end + 1):
            mc2 = mc2.merge(vrow.cells[ci])
        _shade(mc2, TOT_ROW)
        _margins(mc2, 50, 50, 80, 60)
        _write(mc2, labels[gi], bold=True, size=8,
               color=colors[gi], align=WD_ALIGN_PARAGRAPH.CENTER)
        prev_end = end

    _para(doc, "", sa=2)

    # ─────────────────────────────────────────
    #  CERTIFICATION TEXT
    # ─────────────────────────────────────────
    cert = doc.add_paragraph()
    cert.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cert.paragraph_format.space_before = Pt(0)
    cert.paragraph_format.space_after  = Pt(4)
    cr = cert.add_run(
        "I hereby certify on my honor that the above is a true and correct "
        "report of the hours of work performed, record of which was made daily "
        "at the time of arrival at and departure from office."
    )
    cr.font.name   = "Arial"
    cr.font.size   = Pt(8)
    cr.font.italic = True

    # ─────────────────────────────────────────
    #  SIGNATURE BLOCK
    #  Row 0 — empty tall space to write signature
    #  Row 1 — signature line
    #  Row 2 — printed name / verification note
    #  Row 3 — title labels
    # ─────────────────────────────────────────
    st = doc.add_table(rows=4, cols=2)
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    SW = [Inches(3.5), Inches(3.8)]

    def set_row_height(row, height_twips):
        """Force a minimum row height."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height_twips))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    # Row 0 — tall blank space for signature (~1 inch)
    set_row_height(st.rows[0], 1440)  # 1440 twips = 1 inch
    for ci in range(2):
        c = st.rows[0].cells[ci]
        c.width = SW[ci]
        _margins(c, 20, 20, 60, 60)
        _write(c, "", size=8)

    # Row 1 — signature line
    for ci in range(2):
        c = st.rows[1].cells[ci]
        c.width = SW[ci]
        _margins(c, 0, 4, 60, 60)
        _write(c, "_" * 40, size=9)

    # Row 2 — printed name (left) | verification note (right)
    c_left = st.rows[2].cells[0]
    c_left.width = SW[0]
    _margins(c_left, 20, 20, 60, 60)
    _write(c_left, info["name"].upper(), bold=True, size=8.5,
           align=WD_ALIGN_PARAGRAPH.CENTER)

    c_right = st.rows[2].cells[1]
    c_right.width = SW[1]
    _margins(c_right, 20, 20, 60, 60)
    _write(c_right, "Verified as to the prescribed office hours.",
           italic=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 3 — title labels
    c_left3 = st.rows[3].cells[0]
    c_left3.width = SW[0]
    _margins(c_left3, 10, 20, 60, 60)
    _write(c_left3, "Signature of Employee", bold=True, size=8,
           align=WD_ALIGN_PARAGRAPH.CENTER)

    c_right3 = st.rows[3].cells[1]
    c_right3.width = SW[1]
    _margins(c_right3, 10, 20, 60, 60)
    _write(c_right3, "Officer-in-Charge / Supervisor", bold=True, size=8,
           align=WD_ALIGN_PARAGRAPH.CENTER)

    # ─────────────────────────────────────────
    #  FOOTER — pinned to very bottom of page
    #  Uses Word's built-in footer so it never
    #  pushes content onto a second page
    # ─────────────────────────────────────────
    footer_text = (
        f"Official Hours: 8:00 AM – 5:00 PM  |  "
        f"Travel Order / Ticket = 8 hrs credited  |  "
        f"——— = no biometrics log  |  "
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    section = doc.sections[0]
    section.footer_distance = Inches(0.2)
    footer   = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    fr = fp.add_run(footer_text)
    fr.font.name   = "Arial"
    fr.font.size   = Pt(6.5)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(140, 140, 140)

    doc.save(out_path)
    print(f"\n  ✅  DOCX saved: {out_path}")


# =============================================================
#  MAIN
# =============================================================

def main():
    info    = collect_info()
    records = collect_attendance(info)

    print("\n  [PAPER SIZE]")
    print("  [ENTER] Legal (8.5 x 13)   [S] Short Bond (8.5 x 11)")
    ps = input("  > ").strip().upper()
    paper = "short" if ps == "S" else "legal"

    safe  = (info["name"].replace(",", "")
                         .replace(".", "")
                         .replace(" ", "_"))
    fname = f"DTR_{safe}_{info['month_name']}_{info['year']}.docx"

    # Create output folder if it doesn't exist yet
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out   = os.path.join(OUTPUT_DIR, fname)

    print("\n  [GENERATING DOCUMENT]")
    build_docx(info, records, out, paper=paper)

    # Terminal summary
    total_late = sum(r["late"] for r in records)
    total_ut   = sum(r["undertime"] for r in records)
    total_hrs  = round(sum(r["hours"] for r in records), 2)
    present    = sum(1 for r in records
                     if (r["am_in"] or r["is_travel"])
                     and not r["is_weekend"])
    absent     = sum(1 for r in records
                     if r["is_absent"] or r["is_holiday"])

    print(f"""
{"="*60}
  {info["month_name"]} {info["year"]}  —  {info["name"]}
{"="*60}
  Days Present   : {present}
  Days Absent    : {absent}
  Total Hours    : {total_hrs} hrs
  Total Late     : {total_late} min  ({mins_to_readable(total_late)})
  Total Undertime: {total_ut} min  ({mins_to_readable(total_ut)})
{"="*60}
  File: {out}
""")


if __name__ == "__main__":
    main()